import hashlib
import json
import os
import shutil
import sys
import threading
import time
import traceback
from concurrent.futures import ThreadPoolExecutor
from datetime import datetime
import io

import pytesseract
from PyQt5.QtCore import Qt, pyqtSignal, QThreadPool, QRunnable, QObject, QSettings
from PyQt5.QtGui import QDragEnterEvent, QDropEvent, QIcon, QPixmap, QImage
from PyQt5.QtWidgets import (QApplication, QMainWindow, QPushButton, QLabel,
                             QVBoxLayout, QWidget, QFileDialog, QProgressBar, QTextEdit,
                             QMessageBox, QHBoxLayout, QComboBox, QSpinBox, QSlider, QDialog,
                             QDialogButtonBox, QListWidget, QSplitter, QMenu,
                             QSystemTrayIcon, QTabWidget, QStyle, QGroupBox, QLineEdit)
from pdf2image import convert_from_path
from PIL import Image, ImageEnhance


def get_resource_path(relative_path):
    if getattr(sys, 'frozen', False):
        # 如果是打包后的exe
        base_path = sys._MEIPASS
    else:
        # 如果是开发环境
        base_path = os.path.dirname(os.path.abspath(__file__))
    return os.path.join(base_path, relative_path)

def get_tesseract_path():
    # 首先检查打包后的路径
    tesseract_path = get_resource_path(os.path.join('tesseract', 'tesseract.exe'))
    if os.path.exists(tesseract_path):
        return tesseract_path
    
    # 然后检查系统环境变量中的路径
    tesseract_path = shutil.which('tesseract')
    if tesseract_path:
        return tesseract_path
    
    # 最后检查常见安装路径
    common_paths = [
        r'D:\Program Files\Tesseract-OCR\tesseract.exe',
        r'C:\Program Files\Tesseract-OCR\tesseract.exe',
        r'C:\Program Files (x86)\Tesseract-OCR\tesseract.exe'
    ]
    
    for path in common_paths:
        if os.path.exists(path):
            return path
    
    return None

def get_tessdata_path():
    # 首先检查打包后的路径
    tessdata_path = get_resource_path('tessdata')
    if os.path.exists(tessdata_path):
        return tessdata_path
    
    # 检查Tesseract安装目录下的tessdata
    tesseract_path = get_tesseract_path()
    if tesseract_path:
        tessdata_path = os.path.join(os.path.dirname(tesseract_path), 'tessdata')
        if os.path.exists(tessdata_path):
            return tessdata_path
    
    # 检查常见安装路径
    common_paths = [
        r'D:\Program Files\Tesseract-OCR\tessdata',
        r'C:\Program Files\Tesseract-OCR\tessdata',
        r'C:\Program Files (x86)\Tesseract-OCR\tessdata'
    ]
    
    for path in common_paths:
        if os.path.exists(path):
            return path
    
    return None

def get_poppler_path():
    # 首先检查打包后的路径
    poppler_path = get_resource_path('poppler')
    if os.path.exists(poppler_path):
        return poppler_path
    
    # 检查系统环境变量中的路径
    poppler_path = os.environ.get('POPPLER_HOME')
    if poppler_path and os.path.exists(poppler_path):
        return poppler_path
    
    # 检查常见安装路径
    common_paths = [
        r'D:\Program Files\poppler-24.08.0\Library\bin',
        r'C:\Program Files\poppler-24.08.0\Library\bin',
        r'C:\Program Files (x86)\poppler-24.08.0\Library\bin',
        r'D:\Program Files\poppler\bin',
        r'C:\Program Files\poppler\bin',
        r'C:\Program Files (x86)\poppler\bin'
    ]
    
    for path in common_paths:
        if os.path.exists(path):
            return path
    
    return None

def check_dependencies():
    errors = []
    
    # 检查Tesseract
    tesseract_path = get_tesseract_path()
    if not tesseract_path:
        errors.append("找不到Tesseract-OCR")
    else:
        # 检查tessdata目录
        tessdata_path = get_tessdata_path()
        if not tessdata_path:
            errors.append("找不到tessdata目录")
        else:
            # 设置TESSDATA_PREFIX环境变量
            os.environ['TESSDATA_PREFIX'] = tessdata_path
            
            # 检查语言文件
            for lang in ['chi_sim.traineddata', 'eng.traineddata', 'equ.traineddata']:
                lang_path = os.path.join(tessdata_path, lang)
                if not os.path.exists(lang_path):
                    errors.append(f"找不到语言文件: {lang}")
        
        # 检查DLL文件
        tesseract_dir = os.path.dirname(tesseract_path)
        required_dlls = [
            'libtesseract-5.dll',
            'libgcc_s_seh-1.dll',
            'libstdc++-6.dll',
            'libwinpthread-1.dll',
            'zlib1.dll',
            'libpng16-16.dll',
            'libjpeg-8.dll',
            'libtiff-6.dll',
            'libwebp-7.dll'
        ]
        for dll in required_dlls:
            if not os.path.exists(os.path.join(tesseract_dir, dll)):
                errors.append(f"找不到Tesseract组件: {dll}")
    
    # 检查Poppler
    poppler_path = get_poppler_path()
    if not poppler_path:
        errors.append("找不到Poppler")
    else:
        required_dlls = ['pdfinfo.exe', 'pdftoppm.exe']
        for dll in required_dlls:
            if not os.path.exists(os.path.join(poppler_path, dll)):
                errors.append(f"找不到Poppler组件: {dll}")
    
    return errors

class OCRConfigDialog(QDialog):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.setWindowTitle("OCR参数配置")
        self.setModal(True)
        
        layout = QVBoxLayout()
        
        # 语言设置
        language_group = QGroupBox("语言设置")
        language_layout = QVBoxLayout()
        self.language_combo = QComboBox()
        self.language_combo.addItems([
            "中文 (chi_sim)",
            "英文 (eng)",
            "日文 (jpn)",
            "韩文 (kor)",
            "中文+英文 (chi_sim+eng)",
            "自动检测"
        ])
        language_layout.addWidget(QLabel("选择语言:"))
        language_layout.addWidget(self.language_combo)
        language_group.setLayout(language_layout)
        
        # OCR引擎设置
        engine_group = QGroupBox("OCR引擎设置")
        engine_layout = QVBoxLayout()
        
        # DPI设置
        dpi_layout = QHBoxLayout()
        self.dpi_spin = QSpinBox()
        self.dpi_spin.setRange(100, 1200)
        self.dpi_spin.setValue(600)
        self.dpi_spin.setSingleStep(100)
        dpi_layout.addWidget(QLabel("DPI:"))
        dpi_layout.addWidget(self.dpi_spin)
        engine_layout.addLayout(dpi_layout)
        
        # OEM模式
        self.oem_combo = QComboBox()
        self.oem_combo.addItems([
            "0 - 传统模式",
            "1 - LSTM模式",
            "2 - 传统+LSTM",
            "3 - 默认"
        ])
        engine_layout.addWidget(QLabel("OCR引擎模式:"))
        engine_layout.addWidget(self.oem_combo)
        
        # PSM模式
        self.psm_combo = QComboBox()
        self.psm_combo.addItems([
            "0 - 仅方向检测",
            "1 - 自动页面分割+方向检测",
            "2 - 自动页面分割，无方向检测",
            "3 - 全自动页面分割，无方向检测",
            "4 - 假设单列变长文本",
            "5 - 假设统一垂直对齐文本",
            "6 - 假设统一块文本",
            "7 - 假设单行文本",
            "8 - 假设单个单词",
            "9 - 假设单个单词圆形",
            "10 - 假设单个字符",
            "11 - 稀疏文本",
            "12 - 稀疏文本+方向检测",
            "13 - 原始行"
        ])
        engine_layout.addWidget(QLabel("页面分割模式:"))
        engine_layout.addWidget(self.psm_combo)
        
        engine_group.setLayout(engine_layout)
        
        # 图像预处理设置
        preprocess_group = QGroupBox("图像预处理设置")
        preprocess_layout = QVBoxLayout()
        
        # 对比度
        contrast_layout = QHBoxLayout()
        self.contrast_slider = QSlider(Qt.Horizontal)
        self.contrast_slider.setRange(0, 200)
        self.contrast_slider.setValue(100)
        self.contrast_label = QLabel("100%")
        contrast_layout.addWidget(QLabel("对比度:"))
        contrast_layout.addWidget(self.contrast_slider)
        contrast_layout.addWidget(self.contrast_label)
        self.contrast_slider.valueChanged.connect(
            lambda v: self.contrast_label.setText(f"{v}%"))
        
        # 亮度
        brightness_layout = QHBoxLayout()
        self.brightness_slider = QSlider(Qt.Horizontal)
        self.brightness_slider.setRange(0, 200)
        self.brightness_slider.setValue(100)
        self.brightness_label = QLabel("100%")
        brightness_layout.addWidget(QLabel("亮度:"))
        brightness_layout.addWidget(self.brightness_slider)
        brightness_layout.addWidget(self.brightness_label)
        self.brightness_slider.valueChanged.connect(
            lambda v: self.brightness_label.setText(f"{v}%"))
        
        # 锐化
        sharpen_layout = QHBoxLayout()
        self.sharpen_slider = QSlider(Qt.Horizontal)
        self.sharpen_slider.setRange(0, 200)
        self.sharpen_slider.setValue(100)
        self.sharpen_label = QLabel("100%")
        sharpen_layout.addWidget(QLabel("锐化:"))
        sharpen_layout.addWidget(self.sharpen_slider)
        sharpen_layout.addWidget(self.sharpen_label)
        self.sharpen_slider.valueChanged.connect(
            lambda v: self.sharpen_label.setText(f"{v}%"))

        # 识别来源
        source_group = QGroupBox("识别来源")
        source_layout = QVBoxLayout()
        self.source_combo = QComboBox()
        self.source_combo.addItems([
            "本地OCR (Tesseract)",
            "百度OCR (在线)"
        ])
        source_layout.addWidget(QLabel("选择识别方式:"))
        source_layout.addWidget(self.source_combo)
        source_group.setLayout(source_layout)

        # 添加到主布局
        layout.addWidget(source_group)
        
        preprocess_layout.addLayout(contrast_layout)
        preprocess_layout.addLayout(brightness_layout)
        preprocess_layout.addLayout(sharpen_layout)
        preprocess_group.setLayout(preprocess_layout)
        
        # 添加所有组到主布局
        layout.addWidget(language_group)
        layout.addWidget(engine_group)
        layout.addWidget(preprocess_group)
        
        # 添加按钮
        buttons = QDialogButtonBox(
            QDialogButtonBox.Ok | QDialogButtonBox.Cancel,
            Qt.Horizontal, self)
        buttons.accepted.connect(self.accept)
        buttons.rejected.connect(self.reject)
        layout.addWidget(buttons)
        
        self.setLayout(layout)
    
    def get_config(self):
        return {
            'source': self.source_combo.currentText(),
            'language': self.language_combo.currentText(),
            'oem': self.oem_combo.currentIndex(),
            'psm': self.psm_combo.currentIndex(),
            'dpi': self.dpi_spin.value(),
            'contrast': self.contrast_slider.value() / 100.0,
            'brightness': self.brightness_slider.value() / 100.0,
            'sharpen': self.sharpen_slider.value() / 100.0
        }

class BaiduAPISettingsDialog(QDialog):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.setWindowTitle("百度OCR API设置")
        self.setModal(True)
        self.settings = QSettings("PDF_OCR", "BaiduAPI")

        layout = QVBoxLayout()

        # 创建输入框和标签
        for name, label in [
            ("app_id", "App ID"),
            ("api_key", "API Key"),
            ("secret_key", "Secret Key")
        ]:
            # 创建水平布局
            h_layout = QHBoxLayout()
            
            # 添加标签
            h_layout.addWidget(QLabel(label))
            
            # 创建输入框
            input_widget = QLineEdit(self.settings.value(name, ""))
            setattr(self, f"{name}_input", input_widget)
            
            # 如果是密钥，添加显示/隐藏按钮
            if name in ["api_key", "secret_key"]:
                input_widget.setEchoMode(QLineEdit.Password)
                toggle_btn = QPushButton("显示")
                toggle_btn.setCheckable(True)
                toggle_btn.clicked.connect(
                    lambda checked, widget=input_widget: 
                    widget.setEchoMode(QLineEdit.Normal if checked else QLineEdit.Password)
                )
                h_layout.addWidget(toggle_btn)
            
            h_layout.addWidget(input_widget)
            layout.addLayout(h_layout)

        # 添加测试按钮
        test_btn = QPushButton("测试API")
        test_btn.clicked.connect(self.test_api)
        layout.addWidget(test_btn)

        # 添加确定/取消按钮
        button_box = QDialogButtonBox(QDialogButtonBox.Ok | QDialogButtonBox.Cancel)
        button_box.accepted.connect(self.save_settings)
        button_box.rejected.connect(self.reject)
        layout.addWidget(button_box)

        self.setLayout(layout)

    def test_api(self):
        """测试API是否可用"""
        try:
            # 获取当前输入的值
            app_id = self.app_id_input.text().strip()
            api_key = self.api_key_input.text().strip()
            secret_key = self.secret_key_input.text().strip()
            
            if not all([app_id, api_key, secret_key]):
                QMessageBox.warning(self, "错误", "请填写所有字段")
                return
                
            # 创建测试图像
            from PIL import Image, ImageDraw
            img = Image.new('RGB', (100, 100), color='white')
            draw = ImageDraw.Draw(img)
            draw.text((10, 10), "测试", fill='black')
            
            # 调用百度OCR API
            from aip import AipOcr
            client = AipOcr(app_id, api_key, secret_key)
            result = client.basicGeneral(img.tobytes())
            
            if 'words_result' in result:
                QMessageBox.information(self, "成功", "API测试成功！")
            else:
                QMessageBox.warning(self, "错误", f"API返回错误: {result.get('error_msg', '未知错误')}")
                
        except Exception as e:
            QMessageBox.critical(self, "错误", f"测试失败: {str(e)}")

    def save_settings(self):
        """保存设置"""
        # 获取当前输入的值
        app_id = self.app_id_input.text().strip()
        api_key = self.api_key_input.text().strip()
        secret_key = self.secret_key_input.text().strip()
        
        if not all([app_id, api_key, secret_key]):
            QMessageBox.warning(self, "错误", "请填写所有字段")
            return
            
        # 保存设置
        self.settings.setValue("app_id", app_id)
        self.settings.setValue("api_key", api_key)
        self.settings.setValue("secret_key", secret_key)
        self.accept()


class OCRSignals(QObject):
    log = pyqtSignal(str)
    progress = pyqtSignal(int)
    finished = pyqtSignal(str)

class OCRWorker(QRunnable):
    def __init__(self, pdf_path, config, progress_callback, log_callback, finished_callback):
        super().__init__()
        self.pdf_path = pdf_path
        self.config = config
        self.progress_callback = progress_callback
        self.log_callback = log_callback
        self.finished_callback = finished_callback
        self._stop_event = threading.Event()
        self.settings = QSettings("PDF_OCR", "BaiduAPI")
        self.request_queue = []
        self.max_retries = 3
        self.retry_delay = 2  # 秒
        self.cache = {}
        self.stats = {
            'total_pages': 0,
            'processed_pages': 0,
            'total_words': 0,
            'total_lines': 0,
            'total_chars': 0,
            'confidence': 0.0
        }
        self.signals = OCRSignals()

    def _validate_pdf(self):
        """验证PDF文件"""
        try:
            # 检查文件是否存在
            if not os.path.exists(self.pdf_path):
                raise ValueError(f"文件不存在: {self.pdf_path}")
                
            # 检查文件大小
            file_size = os.path.getsize(self.pdf_path)
            if file_size == 0:
                raise ValueError("文件为空")
                
            # 检查文件格式
            with open(self.pdf_path, 'rb') as f:
                header = f.read(4)
                if header != b'%PDF':
                    raise ValueError("不是有效的PDF文件")
                    
            # 检查文件是否可读
            try:
                with open(self.pdf_path, 'rb') as f:
                    f.read(1)
            except IOError:
                raise ValueError("文件无法读取，请检查文件权限")
                
            return True
            
        except Exception as e:
            self.log_callback(f"PDF文件验证失败: {str(e)}")
            raise
            
    def _process_with_baidu(self, img):
        """使用百度OCR处理图像"""
        try:
            # 检查缓存
            cache_key = self._get_cache_key()
            cached_result = self._get_cached_result(cache_key)
            if cached_result:
                if not hasattr(self, '_cache_used'):
                    self.log_callback("使用缓存结果")
                    self._cache_used = True
                return cached_result
                
            # 获取API配置
            app_id = self.settings.value("app_id", "")
            api_key = self.settings.value("api_key", "")
            secret_key = self.settings.value("secret_key", "")
            
            # 检查API配置是否完整
            if not all([app_id, api_key, secret_key]):
                raise ValueError("百度OCR API配置不完整，请先配置API信息")
                
            # 创建客户端
            from aip import AipOcr
            client = AipOcr(app_id, api_key, secret_key)
            
            # 转换为字节
            img_byte = io.BytesIO()
            img.save(img_byte, format='PNG')
            img_byte = img_byte.getvalue()
            
            # 调用API（带重试机制）
            options = {}
            if self.config.get('language') != 'auto':
                options['language_type'] = self.config['language']
                
            for attempt in range(self.max_retries):
                try:
                    # 根据格式选择API
                    if self.config.get('format') == '保留原始格式':
                        result = client.general(img_byte, options)
                    else:
                        result = client.basicGeneral(img_byte, options)
                    
                    if 'error_code' in result:
                        if result['error_code'] == 18:  # QPS超限
                            time.sleep(self.retry_delay)
                            continue
                        raise ValueError(f"百度OCR API错误: {result['error_msg']}")
                        
                    # 提取文本
                    if self.config.get('format') == '保留原始格式':
                        # 使用位置信息保持格式
                        words_result = result.get('words_result', [])
                        lines = []
                        current_line = []
                        last_y = None
                        
                        for item in words_result:
                            location = item.get('location', {})
                            y = location.get('top', 0)
                            
                            if last_y is None:
                                last_y = y
                            elif abs(y - last_y) > 20:  # 如果y坐标差距较大，认为是新行
                                if current_line:
                                    lines.append(' '.join(current_line))
                                    current_line = []
                                last_y = y
                            
                            current_line.append(item['words'])
                        
                        if current_line:
                            lines.append(' '.join(current_line))
                        
                        text = '\n'.join(lines)
                    else:
                        text = '\n'.join([item['words'] for item in result.get('words_result', [])])
                    
                    # 缓存结果
                    self._cache_result(cache_key, text)
                    return text
                    
                except Exception as e:
                    if attempt == self.max_retries - 1:
                        raise
                    time.sleep(self.retry_delay)
                    
        except Exception as e:
            self.log_callback(f"百度OCR处理失败: {str(e)}")
            return None

    def _process_with_tesseract(self, img):
        """使用Tesseract处理图像"""
        try:
            # 检查缓存
            cache_key = self._get_cache_key()
            cached_result = self._get_cached_result(cache_key)
            if cached_result:
                if not hasattr(self, '_cache_used'):
                    self.log_callback("使用缓存结果")
                    self._cache_used = True
                return cached_result
                
            # 设置Tesseract路径
            tesseract_path = get_tesseract_path()
            if tesseract_path:
                pytesseract.pytesseract.tesseract_cmd = tesseract_path
            else:
                raise ValueError("找不到Tesseract-OCR")
                
            # 设置TESSDATA_PREFIX环境变量
            tessdata_path = get_tessdata_path()
            if not tessdata_path:
                raise ValueError("找不到tessdata目录")
            os.environ['TESSDATA_PREFIX'] = tessdata_path
            
            # 获取语言设置
            language = self.config.get('language', 'chi_sim')
            # 处理语言格式
            if isinstance(language, str):
                if '(' in language:
                    language = language.split('(')[-1].strip(')')
                elif language == "自动检测":
                    language = 'auto'
                    
            if language == 'auto':
                # 简单的语言检测
                text = pytesseract.image_to_string(img, lang='chi_sim')
                if len(text.strip()) > 0:
                    language = 'chi_sim'
                else:
                    text = pytesseract.image_to_string(img, lang='eng')
                    if len(text.strip()) > 0:
                        language = 'eng'
                    else:
                        language = 'chi_sim'
                        
            # 检查语言文件
            lang_path = os.path.join(tessdata_path, f"{language}.traineddata")
            if not os.path.exists(lang_path):
                raise ValueError(f"找不到语言文件: {language}")
                
            # 配置Tesseract参数
            custom_config = f'--oem {self.config.get("oem", 1)} --psm {self.config.get("psm", 3)}'
            
            # 预处理图像
            if img.mode != 'L':
                img = img.convert('L')
                
            # 应用对比度和亮度
            contrast = self.config.get('contrast', 1.0)
            brightness = self.config.get('brightness', 1.0)
            
            enhancer = ImageEnhance.Contrast(img)
            img = enhancer.enhance(contrast)
            enhancer = ImageEnhance.Brightness(img)
            img = enhancer.enhance(brightness)
            
            # 获取格式信息
            if self.config.get('format') == '保留原始格式':
                # 使用hOCR输出格式获取位置信息
                hocr = pytesseract.image_to_pdf_or_hocr(
                    img,
                    lang=language,
                    config=custom_config,
                    extension='hocr'
                )
                # 解析hOCR格式
                from bs4 import BeautifulSoup
                soup = BeautifulSoup(hocr, 'html.parser')
                lines = []
                for line in soup.find_all('span', class_='ocr_line'):
                    words = []
                    for word in line.find_all('span', class_='ocrx_word'):
                        words.append(word.get_text().strip())
                    if words:
                        lines.append(' '.join(words))
                text = '\n'.join(lines)
            else:
                # 普通文本输出
                text = pytesseract.image_to_string(
                    img,
                    lang=language,
                    config=custom_config
                )
            
            # 缓存结果
            self._cache_result(cache_key, text)
            return text
            
        except Exception as e:
            self.log_callback(f"Tesseract处理失败: {str(e)}")
            return None

    def _get_cached_result(self, cache_key):
        """获取缓存结果"""
        # 使用配置的缓存目录
        settings = QSettings("PDF_OCR", "CacheSettings")
        cache_dir = settings.value("cache_path", os.path.join(os.path.expanduser('~'), '.pdfocr_cache'))
        os.makedirs(cache_dir, exist_ok=True)
        cache_file = os.path.join(cache_dir, f"{cache_key}.txt")
        
        if os.path.exists(cache_file):
            # 检查缓存是否过期（24小时）
            if os.path.getmtime(cache_file) > (time.time() - 24 * 3600):
                with open(cache_file, 'r', encoding='utf-8') as f:
                    return f.read()
        return None

    def _cache_result(self, cache_key, result):
        """缓存结果"""
        # 使用配置的缓存目录
        settings = QSettings("PDF_OCR", "CacheSettings")
        cache_dir = settings.value("cache_path", os.path.join(os.path.expanduser('~'), '.pdfocr_cache'))
        os.makedirs(cache_dir, exist_ok=True)
        cache_file = os.path.join(cache_dir, f"{cache_key}.txt")
        
        with open(cache_file, 'w', encoding='utf-8') as f:
            f.write(result)

    def _get_cache_key(self):
        """生成缓存键"""
        # 使用文件路径和修改时间生成缓存键
        file_mtime = os.path.getmtime(self.pdf_path)
        file_info = f"{self.pdf_path}_{file_mtime}"
        config_hash = hashlib.md5(json.dumps(self.config, sort_keys=True).encode()).hexdigest()
        return f"{hashlib.md5(file_info.encode()).hexdigest()}_{config_hash}"

    def _update_stats(self, text):
        """更新统计信息"""
        lines = text.split('\n')
        words = sum(len(line.split()) for line in lines)
        chars = sum(len(line) for line in lines)
        
        self.stats['total_lines'] += len(lines)
        self.stats['total_words'] += words
        self.stats['total_chars'] += chars
        self.stats['processed_pages'] += 1
        
        # 计算平均置信度
        if self.stats['processed_pages'] > 0:
            self.stats['confidence'] = self.stats['confidence'] / self.stats['processed_pages']
            
        return {
            'pages': f"{self.stats['processed_pages']}/{self.stats['total_pages']}",
            'lines': self.stats['total_lines'],
            'words': self.stats['total_words'],
            'chars': self.stats['total_chars'],
            'confidence': f"{self.stats['confidence']:.1f}%"
        }

    def _convert_image_to_qimage(self, img):
        """将PIL图像转换为QImage"""
        try:
            if isinstance(img, Image.Image):
                # 确保图像是RGB模式
                if img.mode != 'RGB':
                    img = img.convert('RGB')
                # 转换为字节数据
                data = img.tobytes('raw', 'RGB')
                # 创建QImage
                qimage = QImage(data, img.size[0], img.size[1], QImage.Format_RGB888)
                return qimage
        except Exception as e:
            self.log_callback(f"图像转换错误: {str(e)}")
        return QImage()

    def run(self):
        """处理PDF文件"""
        try:
            # 重置缓存使用标记
            self._cache_used = False
            
            # 验证PDF文件
            self._validate_pdf()
            
            # 检查API配置
            if self.config.get('source') == '百度OCR (在线)':
                app_id = self.settings.value("app_id", "")
                api_key = self.settings.value("api_key", "")
                secret_key = self.settings.value("secret_key", "")
                
                if not all([app_id, api_key, secret_key]):
                    raise ValueError("百度OCR API配置不完整，请先配置API信息")
                    
            # 获取Poppler路径
            poppler_path = get_poppler_path()
            if not poppler_path:
                raise ValueError("找不到Poppler，请确保程序完整性")
                
            # 转换PDF为图像
            try:
                self.log_callback("开始转换PDF...")
                self.log_callback(f"使用DPI: {self.config.get('dpi', 300)}")
                self.log_callback(f"Poppler路径: {poppler_path}")
                
                # 检查Poppler工具
                pdfinfo_path = os.path.join(poppler_path, 'pdfinfo.exe')
                if not os.path.exists(pdfinfo_path):
                    raise ValueError(f"找不到pdfinfo工具: {pdfinfo_path}")
                    
                images = convert_from_path(
                    self.pdf_path,
                    poppler_path=poppler_path,
                    dpi=self.config.get('dpi', 300),
                    thread_count=1  # 使用单线程避免并发问题
                )
                self.log_callback(f"PDF转换完成，共{len(images)}页")
            except Exception as e:
                self.log_callback(f"PDF转换错误: {str(e)}")
                self.log_callback(traceback.format_exc())
                raise Exception(f"PDF转换失败: {str(e)}")
                
            total_pages = len(images)
            
            # 设置总页数
            self.stats['total_pages'] = total_pages
            
            # 创建输出目录
            output_dir = os.path.dirname(self.pdf_path)
            if not output_dir:
                output_dir = os.getcwd()
            os.makedirs(output_dir, exist_ok=True)
            
            # 处理每一页
            output_path = os.path.splitext(self.pdf_path)[0] + '_ocr.txt'
            result_text = ""
            with open(output_path, 'w', encoding='utf-8') as f:
                for i, img in enumerate(images):
                    try:
                        # 更新进度
                        self.progress_callback(int((i + 1) / total_pages * 100))
                        
                        # 处理页面
                        if self.config.get('source') == '百度OCR (在线)':
                            text = self._process_with_baidu(img)
                        else:
                            text = self._process_with_tesseract(img)
                            
                        if text:
                            page_text = f"=== 第 {i+1} 页 ===\n{text}\n\n"
                            f.write(page_text)
                            result_text += page_text
                            
                            # 更新统计信息
                            stats = self._update_stats(text)
                            
                    except Exception as e:
                        self.log_callback(f"处理第 {i+1} 页时出错: {str(e)}")
                        continue
                        
            self.finished_callback(result_text)  # 传递结果文本
            
        except Exception as e:
            self.finished_callback(f"处理PDF时出错: {str(e)}")  # 传递错误信息
    
    def stop(self):
        self._stop_event.set()

class CacheSettingsDialog(QDialog):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.setWindowTitle("缓存设置")
        self.setModal(True)
        
        # 加载设置
        self.settings = QSettings("PDF_OCR", "CacheSettings")
        self.default_cache_path = os.path.join(os.path.expanduser('~'), '.pdfocr_cache')
        self.cache_path = self.settings.value("cache_path", self.default_cache_path)
        
        layout = QVBoxLayout()
        
        # 缓存信息组
        info_group = QGroupBox("缓存信息")
        info_layout = QVBoxLayout()
        
        # 缓存路径设置
        path_layout = QHBoxLayout()
        self.cache_path_label = QLabel(self.cache_path)
        self.change_path_button = QPushButton("更改路径")
        self.change_path_button.clicked.connect(self.change_cache_path)
        self.reset_path_button = QPushButton("恢复默认")
        self.reset_path_button.clicked.connect(self.reset_cache_path)
        path_layout.addWidget(QLabel("缓存路径:"))
        path_layout.addWidget(self.cache_path_label)
        path_layout.addWidget(self.change_path_button)
        path_layout.addWidget(self.reset_path_button)
        info_layout.addLayout(path_layout)
        
        # 缓存大小
        self.cache_size_label = QLabel()
        info_layout.addWidget(QLabel("缓存大小:"))
        info_layout.addWidget(self.cache_size_label)
        
        info_group.setLayout(info_layout)
        
        # 操作按钮组
        button_group = QGroupBox("操作")
        button_layout = QVBoxLayout()
        
        # 打开缓存文件夹按钮
        self.open_folder_button = QPushButton("打开缓存文件夹")
        self.open_folder_button.clicked.connect(self.open_cache_folder)
        
        # 清除缓存按钮
        self.clear_cache_button = QPushButton("清除缓存")
        self.clear_cache_button.clicked.connect(self.clear_cache)
        
        button_layout.addWidget(self.open_folder_button)
        button_layout.addWidget(self.clear_cache_button)
        button_group.setLayout(button_layout)
        
        # 添加所有组到主布局
        layout.addWidget(info_group)
        layout.addWidget(button_group)
        
        # 添加关闭按钮
        close_button = QPushButton("关闭")
        close_button.clicked.connect(self.accept)
        layout.addWidget(close_button)
        
        self.setLayout(layout)
        
        # 更新缓存信息
        self.update_cache_info()
    
    def change_cache_path(self):
        """更改缓存路径"""
        new_path = QFileDialog.getExistingDirectory(
            self,
            "选择缓存目录",
            self.cache_path,
            QFileDialog.ShowDirsOnly | QFileDialog.DontResolveSymlinks
        )
        
        if new_path:
            # 检查新路径是否可写
            test_file = os.path.join(new_path, '.test')
            try:
                with open(test_file, 'w') as f:
                    f.write('test')
                os.remove(test_file)
            except Exception as e:
                QMessageBox.warning(
                    self,
                    "错误",
                    f"无法写入选择的目录：\n{str(e)}\n请选择其他目录。"
                )
                return
            
            # 如果旧路径存在，询问是否移动现有缓存
            if os.path.exists(self.cache_path) and os.listdir(self.cache_path):
                reply = QMessageBox.question(
                    self,
                    "移动缓存",
                    "是否将现有缓存文件移动到新位置？",
                    QMessageBox.Yes | QMessageBox.No,
                    QMessageBox.Yes
                )
                
                if reply == QMessageBox.Yes:
                    try:
                        # 创建新目录
                        os.makedirs(new_path, exist_ok=True)
                        
                        # 移动文件
                        for file in os.listdir(self.cache_path):
                            src = os.path.join(self.cache_path, file)
                            dst = os.path.join(new_path, file)
                            shutil.move(src, dst)
                        
                        # 删除旧目录
                        if not os.listdir(self.cache_path):
                            os.rmdir(self.cache_path)
                    except Exception as e:
                        QMessageBox.warning(
                            self,
                            "错误",
                            f"移动缓存文件时出错：\n{str(e)}"
                        )
                        return
            
            # 更新路径
            self.cache_path = new_path
            self.cache_path_label.setText(self.cache_path)
            
            # 保存设置
            self.settings.setValue("cache_path", self.cache_path)
            
            # 更新缓存信息
            self.update_cache_info()
    
    def reset_cache_path(self):
        """恢复默认缓存路径"""
        reply = QMessageBox.question(
            self,
            "确认恢复默认",
            "确定要恢复默认缓存路径吗？\n当前缓存文件将不会被移动。",
            QMessageBox.Yes | QMessageBox.No,
            QMessageBox.No
        )
        
        if reply == QMessageBox.Yes:
            # 更新路径
            self.cache_path = self.default_cache_path
            self.cache_path_label.setText(self.cache_path)
            
            # 保存设置
            self.settings.setValue("cache_path", self.cache_path)
            
            # 更新缓存信息
            self.update_cache_info()
            
            QMessageBox.information(
                self,
                "恢复完成",
                "已恢复默认缓存路径"
            )
    
    def update_cache_info(self):
        """更新缓存信息显示"""
        if os.path.exists(self.cache_path):
            # 计算缓存大小
            total_size = 0
            for root, dirs, files in os.walk(self.cache_path):
                for file in files:
                    file_path = os.path.join(root, file)
                    total_size += os.path.getsize(file_path)
            
            # 格式化显示大小
            if total_size < 1024:
                size_str = f"{total_size} 字节"
            elif total_size < 1024 * 1024:
                size_str = f"{total_size/1024:.2f} KB"
            else:
                size_str = f"{total_size/1024/1024:.2f} MB"
            
            self.cache_size_label.setText(size_str)
        else:
            self.cache_size_label.setText("0 字节")
    
    def open_cache_folder(self):
        """打开缓存文件夹"""
        if os.path.exists(self.cache_path):
            os.startfile(self.cache_path)
        else:
            QMessageBox.information(self, "提示", "暂无缓存文件夹")
    
    def clear_cache(self):
        """清除所有缓存文件"""
        reply = QMessageBox.question(
            self,
            "确认清除缓存",
            "确定要清除所有OCR缓存文件吗？\n这将删除所有已保存的识别结果。",
            QMessageBox.Yes | QMessageBox.No,
            QMessageBox.No
        )
        
        if reply == QMessageBox.Yes:
            try:
                if os.path.exists(self.cache_path):
                    # 删除所有缓存文件
                    deleted_count = 0
                    for file in os.listdir(self.cache_path):
                        file_path = os.path.join(self.cache_path, file)
                        if os.path.isfile(file_path):
                            os.remove(file_path)
                            deleted_count += 1
                    
                    QMessageBox.information(
                        self,
                        "清除完成",
                        f"已清除 {deleted_count} 个缓存文件"
                    )
                    
                    # 更新缓存信息
                    self.update_cache_info()
                else:
                    QMessageBox.information(self, "提示", "暂无缓存文件")
            except Exception as e:
                QMessageBox.warning(
                    self,
                    "清除失败",
                    f"清除缓存时发生错误：\n{str(e)}"
                )

class MainWindow(QMainWindow):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("PDF OCR识别工具")
        # 设置主窗口图标
        icon_path = get_resource_path('icon.ico')
        self.setWindowIcon(QIcon(icon_path))
        
        # 创建主窗口部件
        main_widget = QWidget()
        self.setCentralWidget(main_widget)
        main_layout = QHBoxLayout()
        
        # 添加线程池
        self.thread_pool = QThreadPool()
        self.thread_pool.setMaxThreadCount(4)
        
        # 创建信号对象
        self.signals = OCRSignals()
        
        # 创建左侧面板（历史记录和最近文件）
        left_panel = QWidget()
        left_layout = QVBoxLayout()
        
        # 创建选项卡
        self.tab_widget = QTabWidget()
        
        # 历史记录标签页
        history_tab = QWidget()
        history_layout = QVBoxLayout()
        self.history_list = QListWidget()
        self.history_list.itemClicked.connect(self.load_history_item)
        self.history_list.setMinimumWidth(300)
        self.history_list.setMaximumWidth(400)
        self.history_list.setContextMenuPolicy(Qt.CustomContextMenu)
        self.history_list.customContextMenuRequested.connect(self.show_history_context_menu)
        history_layout.addWidget(self.history_list)
        history_tab.setLayout(history_layout)
        
        # 最近文件标签页
        recent_tab = QWidget()
        recent_layout = QVBoxLayout()
        self.recent_list = QListWidget()
        self.recent_list.itemClicked.connect(self.load_recent_file)
        self.recent_list.setMinimumWidth(300)
        self.recent_list.setMaximumWidth(400)
        self.recent_list.setContextMenuPolicy(Qt.CustomContextMenu)
        self.recent_list.customContextMenuRequested.connect(self.show_recent_context_menu)
        recent_layout.addWidget(self.recent_list)
        recent_tab.setLayout(recent_layout)
        
        # 添加选项卡
        self.tab_widget.addTab(history_tab, "历史记录")
        self.tab_widget.addTab(recent_tab, "最近文件")
        
        left_layout.addWidget(self.tab_widget)
        left_panel.setLayout(left_layout)
        
        # 创建右侧面板（主功能）
        right_panel = QWidget()
        right_layout = QVBoxLayout()
        
        # 创建主题切换按钮
        theme_layout = QHBoxLayout()
        self.theme_button = QPushButton("切换主题")
        self.theme_button.clicked.connect(self.toggle_theme)
        theme_layout.addWidget(self.theme_button)
        theme_layout.addStretch()
        right_layout.addLayout(theme_layout)
        
        # 创建水平布局用于按钮
        button_layout = QHBoxLayout()
        
        # 创建按钮
        self.select_button = QPushButton("选择PDF文件")
        self.select_button.clicked.connect(self.select_pdf)
        
        self.select_multiple_button = QPushButton("批量选择PDF")
        self.select_multiple_button.clicked.connect(self.select_multiple_pdf)
        
        self.copy_button = QPushButton("复制文本")
        self.copy_button.clicked.connect(self.copy_text)
        self.copy_button.setEnabled(False)
        
        self.export_button = QPushButton("导出文本")
        self.export_button.clicked.connect(self.export_text)
        self.export_button.setEnabled(False)
        
        self.export_word_button = QPushButton("导出Word")
        self.export_word_button.clicked.connect(self.export_word)
        self.export_word_button.setEnabled(False)
        
        # 添加配置按钮
        self.config_button = QPushButton("OCR配置")
        self.config_button.clicked.connect(self.show_config_dialog)
        
        # 添加取消按钮
        self.cancel_button = QPushButton("取消")
        self.cancel_button.clicked.connect(self.cancel_ocr)
        self.cancel_button.setEnabled(False)
        
        # 添加缓存设置按钮
        self.cache_settings_button = QPushButton("缓存设置")
        self.cache_settings_button.clicked.connect(self.show_cache_settings)

        self.api_settings_button = QPushButton("百度API设置")
        self.api_settings_button.clicked.connect(self.show_baidu_api_settings)
        
        # 添加批量导出按钮
        self.batch_export_button = QPushButton("批量导出")
        self.batch_export_button.clicked.connect(self.batch_export)
        button_layout.addWidget(self.batch_export_button)
        
        # 添加校对按钮
        self.proofread_button = QPushButton("文本校对")
        self.proofread_button.clicked.connect(self.proofread_text)
        self.proofread_button.setEnabled(False)
        button_layout.addWidget(self.proofread_button)
        
        # 添加批量处理进度条
        self.batch_progress = QProgressBar()
        self.batch_progress.setVisible(False)
        right_layout.addWidget(self.batch_progress)
        
        # 添加按钮到水平布局
        button_layout.addWidget(self.select_button)
        button_layout.addWidget(self.select_multiple_button)
        button_layout.addWidget(self.copy_button)
        button_layout.addWidget(self.export_button)
        button_layout.addWidget(self.export_word_button)
        button_layout.addWidget(self.config_button)
        button_layout.addWidget(self.cancel_button)
        button_layout.addWidget(self.cache_settings_button)
        button_layout.addWidget(self.api_settings_button)
        
        # 创建设置区域
        settings_layout = QHBoxLayout()
        
        # 添加格式选项
        self.format_combo = QComboBox()
        self.format_combo.addItems(["保留原始格式", "纯文本", "Markdown"])
        settings_layout.addWidget(QLabel("输出格式:"))
        settings_layout.addWidget(self.format_combo)
        
        # 添加字体大小设置
        self.font_size_spin = QSpinBox()
        self.font_size_spin.setRange(8, 24)
        self.font_size_spin.setValue(12)
        self.font_size_spin.setSingleStep(1)
        self.font_size_spin.valueChanged.connect(self.update_font_size)
        settings_layout.addWidget(QLabel("字体大小:"))
        settings_layout.addWidget(self.font_size_spin)
        
        self.progress_bar = QProgressBar()
        self.progress_bar.setVisible(False)
        
        # 创建日志显示区域
        self.log_text = QTextEdit()
        self.log_text.setReadOnly(True)
        self.log_text.setMinimumHeight(200)
        self.log_text.setMaximumHeight(300)
        self.log_text.setStyleSheet("""
            QTextEdit {
                background-color: #f0f0f0;
                font-family: 'Consolas', 'Courier New', monospace;
            }
        """)
        
        # 创建结果显示区域
        self.result_text = QTextEdit()
        self.result_text.setReadOnly(False)
        self.result_text.setStyleSheet("""
            QTextEdit {
                font-family: 'Consolas', 'Courier New', monospace;
            }
        """)
        
        # 添加统计信息显示
        self.stats_label = QLabel()
        self.stats_label.setStyleSheet("color: #666666;")
        
        # 添加部件到右侧布局
        right_layout.addLayout(button_layout)
        right_layout.addLayout(settings_layout)
        right_layout.addWidget(self.progress_bar)
        right_layout.addWidget(self.log_text)
        right_layout.addWidget(self.result_text)
        right_layout.addWidget(self.proofread_button)
        right_layout.addWidget(self.stats_label)
        right_layout.addWidget(self.batch_progress)
        
        right_panel.setLayout(right_layout)
        
        # 创建分割器
        splitter = QSplitter(Qt.Horizontal)
        splitter.addWidget(left_panel)
        splitter.addWidget(right_panel)
        splitter.setStretchFactor(0, 1)
        splitter.setStretchFactor(1, 4)
        
        main_layout.addWidget(splitter)
        main_widget.setLayout(main_layout)
        
        self.ocr_thread = None
        self.current_theme = "浅色"
        self.history = []
        self.recent_files = []
        
        # 加载设置
        self.settings = QSettings("PDF_OCR", "Settings")
        self.load_settings()
        
        # 加载OCR配置
        self.ocr_config = self.settings.value("ocr_config", {
            'language': '中文 (chi_sim)',
            'oem': 1,  # 使用LSTM模式，速度更快
            'psm': 3,  # 全自动页面分割，无方向检测
            'dpi': 300,  # 降低DPI以提高速度
            'contrast': 1.0,
            'brightness': 1.0,
            'sharpen': 1.0
        })
        
        self.load_history()
        self.load_recent_files()
        
        # 启用拖放
        self.setAcceptDrops(True)
        
        # 检查依赖
        errors = check_dependencies()
        if errors:
            error_msg = "程序初始化失败:\n" + "\n".join(errors)
            QMessageBox.critical(self, "错误", error_msg)
            self.select_button.setEnabled(False)
            self.select_multiple_button.setEnabled(False)
            self.result_text.setText(error_msg)
        
        # 连接信号
        self.signals.log.connect(self._update_log)
        self.signals.progress.connect(self._update_progress)
        self.signals.finished.connect(self._ocr_finished)
        
        # 创建系统托盘图标
        self.create_tray_icon()
        
        # 添加校对窗口
        self.proofread_dialog = None
        
        # 添加状态栏
        self.statusBar().showMessage("就绪")
    
    def load_settings(self):
        # 加载窗口大小和位置
        geometry = self.settings.value("geometry")
        if geometry:
            self.restoreGeometry(geometry)
        else:
            self.setGeometry(100, 100, 1200, 800)
        
        # 加载主题设置
        self.current_theme = self.settings.value("theme", "浅色")
        self.apply_theme()
        
        # 加载字体大小设置
        font_size = self.settings.value("font_size", 12, type=int)
        self.font_size_spin.setValue(font_size)
        self.update_font_size(font_size)
    
    def save_settings(self):
        # 保存窗口大小和位置
        self.settings.setValue("geometry", self.saveGeometry())
        # 保存主题设置
        self.settings.setValue("theme", self.current_theme)
    
    def apply_theme(self):
        if self.current_theme == "深色":
            self.setStyleSheet("""
                QMainWindow, QWidget {
                    background-color: #2b2b2b;
                    color: #ffffff;
                }
                QPushButton {
                    background-color: #3c3f41;
                    color: #ffffff;
                    border: 1px solid #555555;
                }
                QPushButton:hover {
                    background-color: #4c4f51;
                }
                QTextEdit {
                    background-color: #323232;
                    color: #ffffff;
                }
                QComboBox, QSpinBox {
                    background-color: #3c3f41;
                    color: #ffffff;
                }
                QListWidget {
                    background-color: #323232;
                    color: #ffffff;
                }
                QTabWidget::pane {
                    border: 1px solid #555555;
                    background-color: #2b2b2b;
                }
                QTabBar::tab {
                    background-color: #3c3f41;
                    color: #ffffff;
                    border: 1px solid #555555;
                    padding: 8px;
                }
                QTabBar::tab:selected {
                    background-color: #4c4f51;
                }
            """)
        else:
            self.setStyleSheet("")
    
    def toggle_theme(self):
        self.current_theme = "深色" if self.current_theme == "浅色" else "浅色"
        self.apply_theme()
        self.save_settings()
    
    def create_tray_icon(self):
        self.tray_icon = QSystemTrayIcon(self)
        # 使用icon.ico作为托盘图标
        icon_path = get_resource_path('icon.ico')
        self.tray_icon.setIcon(QIcon(icon_path))
        
        # 创建托盘菜单
        tray_menu = QMenu()
        show_action = tray_menu.addAction("显示")
        show_action.triggered.connect(self.show)
        quit_action = tray_menu.addAction("退出")
        quit_action.triggered.connect(QApplication.quit)
        
        self.tray_icon.setContextMenu(tray_menu)
        self.tray_icon.show()
        
        # 添加双击显示功能
        self.tray_icon.activated.connect(self.tray_icon_activated)
    
    def tray_icon_activated(self, reason):
        if reason == QSystemTrayIcon.DoubleClick:
            self.show()
            self.activateWindow()
    
    def closeEvent(self, event):
        # 保存设置
        self.save_settings()
        # 最小化到托盘
        if self.tray_icon.isVisible():
            self.hide()
            event.ignore()
        else:
            event.accept()
    
    def dragEnterEvent(self, event: QDragEnterEvent):
        if event.mimeData().hasUrls():
            event.acceptProposedAction()
    
    def dropEvent(self, event: QDropEvent):
        files = [url.toLocalFile() for url in event.mimeData().urls()]
        pdf_files = [f for f in files if f.lower().endswith('.pdf')]
        if pdf_files:
            if len(pdf_files) == 1:
                self.start_ocr(pdf_files[0])
            else:
                self.batch_process_pdfs(pdf_files)
    
    def load_history(self):
        try:
            with open('history.json', 'r', encoding='utf-8') as f:
                self.history = json.load(f)
                self.update_history_list()
        except FileNotFoundError:
            self.history = []
    
    def save_history(self):
        with open('history.json', 'w', encoding='utf-8') as f:
            json.dump(self.history, f, ensure_ascii=False, indent=2)
    
    def update_history_list(self):
        self.history_list.clear()
        for item in reversed(self.history):
            # 获取文件大小
            file_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), item['filename'])
            file_size = os.path.getsize(file_path) if os.path.exists(file_path) else 0
            file_size_str = f"{file_size/1024:.1f}KB" if file_size < 1024*1024 else f"{file_size/1024/1024:.1f}MB"
            
            # 获取文本统计信息
            text = item['result']
            lines = len(text.split('\n'))
            words = sum(len(line.split()) for line in text.split('\n'))
            chars = len(text)
            
            # 格式化显示
            display_text = (
                f"📄 {item['filename']}\n"
                f"⏰ {item['time']}\n"
                f"📊 {file_size_str} | {lines}行 | {words}字 | {chars}字符"
            )
            self.history_list.addItem(display_text)
    
    def add_to_history(self, filename, result):
        """添加到历史记录"""
        # 检查是否已存在相同文件的历史记录
        for item in self.history:
            if item['filename'] == filename:
                # 更新现有记录的时间和结果
                item['time'] = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
                item['result'] = result
                self.save_history()
                self.update_history_list()
                return
                
        # 如果是新文件，添加新记录
        history_item = {
            'time': datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
            'filename': filename,
            'result': result
        }
        self.history.append(history_item)
        self.save_history()
        self.update_history_list()
    
    def load_history_item(self, item):
        index = self.history_list.row(item)
        history_item = self.history[-(index + 1)]
        self.result_text.setText(history_item['result'])
        self.copy_button.setEnabled(True)
        self.export_button.setEnabled(True)
        self.export_word_button.setEnabled(True)
        self.proofread_button.setEnabled(True)
    
    def proofread_text(self):
        """文本校对功能"""
        if not self.proofread_dialog:
            self.proofread_dialog = ProofreadDialog(self)
            
        self.proofread_dialog.set_text(self.result_text.toPlainText())
        if self.proofread_dialog.exec_() == QDialog.Accepted:
            self.result_text.setText(self.proofread_dialog.get_text())
    
    def export_word(self):
        try:
            from docx import Document
            file_path, _ = QFileDialog.getSaveFileName(
                self,
                "导出Word文档",
                "",
                "Word文档 (*.docx)"
            )
            
            if file_path:
                if not file_path.lower().endswith('.docx'):
                    file_path += '.docx'
                
                doc = Document()
                doc.add_paragraph(self.result_text.toPlainText())
                doc.save(file_path)
                QMessageBox.information(self, "提示", f"Word文档已成功导出到：\n{file_path}")
        except ImportError:
            QMessageBox.warning(self, "警告", "请先安装python-docx库：\npip install python-docx")
    
    def show_config_dialog(self):
        dialog = OCRConfigDialog(self)
        # 设置当前配置
        dialog.language_combo.setCurrentText(self.ocr_config.get('language', '中文 (chi_sim)'))
        dialog.oem_combo.setCurrentIndex(self.ocr_config.get('oem', 1))
        dialog.psm_combo.setCurrentIndex(self.ocr_config.get('psm', 3))
        dialog.dpi_spin.setValue(self.ocr_config.get('dpi', 300))
        dialog.contrast_slider.setValue(int(self.ocr_config.get('contrast', 1.0) * 100))
        dialog.brightness_slider.setValue(int(self.ocr_config.get('brightness', 1.0) * 100))
        dialog.sharpen_slider.setValue(int(self.ocr_config.get('sharpen', 1.0) * 100))
        # 设置识别来源
        source = self.ocr_config.get('source', '本地OCR (Tesseract)')
        dialog.source_combo.setCurrentText(source)
        
        if dialog.exec_() == QDialog.Accepted:
            self.ocr_config = dialog.get_config()
            # 保存OCR配置
            self.settings.setValue("ocr_config", self.ocr_config)
            self.log_text.append("OCR配置已更新并保存")
    
    def update_stats(self, text):
        # 更新统计信息
        lines = text.split('\n')
        words = sum(len(line.split()) for line in lines)
        chars = sum(len(line) for line in lines)
        self.stats_label.setText(
            f"行数: {len(lines)} | 字数: {words} | 字符数: {chars}"
        )
    
    def _update_log(self, message):
        self.log_text.append(message)
        # 滚动到底部
        self.log_text.verticalScrollBar().setValue(
            self.log_text.verticalScrollBar().maximum()
        )
        # 确保日志窗口可见
        self.log_text.setVisible(True)
    
    def _update_progress(self, value):
        self.progress_bar.setValue(value)
    
    def _ocr_finished(self, result):
        self.result_text.setText(result)
        self.select_button.setEnabled(True)
        self.select_multiple_button.setEnabled(True)
        self.copy_button.setEnabled(True)
        self.export_button.setEnabled(True)
        self.export_word_button.setEnabled(True)
        self.proofread_button.setEnabled(True)
        self.config_button.setEnabled(True)
        self.cancel_button.setEnabled(False)
        self.progress_bar.setVisible(False)
        
        # 更新统计信息
        self.update_stats(result)
        
        # 添加到历史记录
        if hasattr(self, 'current_pdf_path'):
            self.add_to_history(self.current_pdf_path, result)
    
    def change_theme(self, theme):
        if theme == "深色" and self.current_theme == "浅色":
            self.setStyleSheet("""
                QMainWindow, QWidget {
                    background-color: #2b2b2b;
                    color: #ffffff;
                }
                QPushButton {
                    background-color: #3c3f41;
                    color: #ffffff;
                    border: 1px solid #555555;
                }
                QPushButton:hover {
                    background-color: #4c4f51;
                }
                QTextEdit {
                    background-color: #323232;
                    color: #ffffff;
                }
                QComboBox, QSpinBox {
                    background-color: #3c3f41;
                    color: #ffffff;
                }
            """)
            self.current_theme = "深色"
        elif theme == "浅色" and self.current_theme == "深色":
            self.setStyleSheet("")
            self.current_theme = "浅色"
    
    def select_multiple_pdf(self):
        file_paths, _ = QFileDialog.getOpenFileNames(
            self,
            "选择多个PDF文件",
            "",
            "PDF文件 (*.pdf)"
        )
        
        if file_paths:
            self.batch_process_pdfs(file_paths)
    
    def batch_process_pdfs(self, file_paths):
        total_files = len(file_paths)
        self.progress_bar.setMaximum(total_files)
        self.progress_bar.setValue(0)
        self.progress_bar.setVisible(True)
        
        for i, file_path in enumerate(file_paths):
            self.log_text.append(f"正在处理文件 {i+1}/{total_files}: {os.path.basename(file_path)}")
            self.start_ocr(file_path)
            self.progress_bar.setValue(i + 1)
        
        self.progress_bar.setVisible(False)
        QMessageBox.information(self, "完成", f"已处理 {total_files} 个文件")
    
    def select_pdf(self):
        file_path, _ = QFileDialog.getOpenFileName(
            self,
            "选择PDF文件",
            "",
            "PDF文件 (*.pdf)"
        )
        
        if file_path:
            self.start_ocr(file_path)
    
    def copy_text(self):
        clipboard = QApplication.clipboard()
        clipboard.setText(self.result_text.toPlainText())
        QMessageBox.information(self, "提示", "文本已复制到剪贴板")
    
    def export_text(self):
        file_path, _ = QFileDialog.getSaveFileName(
            self,
            "导出文本文件",
            "",
            "文本文件 (*.txt)"
        )
        
        if file_path:
            try:
                # 确保文件路径有.txt扩展名
                if not file_path.lower().endswith('.txt'):
                    file_path += '.txt'
                
                # 获取文本内容
                text_content = self.result_text.toPlainText()
                
                # 检查文本内容是否为空
                if not text_content.strip():
                    QMessageBox.warning(self, "警告", "没有可导出的文本内容")
                    return
                
                # 写入文件
                with open(file_path, 'w', encoding='utf-8') as f:
                    f.write(text_content)
                
                QMessageBox.information(self, "提示", f"文本已成功导出到：\n{file_path}")
            except Exception as e:
                QMessageBox.critical(self, "错误", f"导出失败: {str(e)}\n请检查文件路径是否有写入权限。")
    
    def start_ocr(self, pdf_path):
        self.select_button.setEnabled(False)
        self.select_multiple_button.setEnabled(False)
        self.copy_button.setEnabled(False)
        self.export_button.setEnabled(False)
        self.export_word_button.setEnabled(False)
        self.proofread_button.setEnabled(False)
        self.config_button.setEnabled(False)
        self.cancel_button.setEnabled(True)
        self.progress_bar.setVisible(True)
        self.progress_bar.setValue(0)
        self.result_text.clear()
        self.log_text.clear()
        
        self.current_pdf_path = pdf_path
        # 添加到最近文件列表
        self.add_to_recent(pdf_path)
        
        self.current_worker = OCRWorker(
            pdf_path,
            self.ocr_config,
            self.signals.progress.emit,
            self.signals.log.emit,
            self.signals.finished.emit
        )
        self.thread_pool.start(self.current_worker)
    
    def cancel_ocr(self):
        if hasattr(self, 'current_worker'):
            self.current_worker.stop()
            self.cancel_button.setEnabled(False)
            self.select_button.setEnabled(True)
            self.select_multiple_button.setEnabled(True)
            self.config_button.setEnabled(True)
            self.progress_bar.setVisible(False)
            self.log_text.append("处理已取消")
    
    def show_history_context_menu(self, position):
        # 创建右键菜单
        menu = QMenu()
        delete_action = menu.addAction("删除")
        clear_all_action = menu.addAction("清空所有")
        
        # 获取右键点击的项目
        item = self.history_list.itemAt(position)
        if item:
            # 执行删除操作
            action = menu.exec_(self.history_list.mapToGlobal(position))
            if action == delete_action:
                self.delete_history_item(item)
            elif action == clear_all_action:
                self.clear_all_history()
        else:
            # 如果没有选中项目，只显示清空所有选项
            action = menu.exec_(self.history_list.mapToGlobal(position))
            if action == clear_all_action:
                self.clear_all_history()
    
    def delete_history_item(self, item):
        # 获取要删除的项目的索引
        index = self.history_list.row(item)
        # 从历史记录中删除
        del self.history[-(index + 1)]
        # 保存更新后的历史记录
        self.save_history()
        # 更新历史记录列表显示
        self.update_history_list()
    
    def clear_all_history(self):
        reply = QMessageBox.question(self, '确认', '确定要清空所有历史记录吗？',
                                   QMessageBox.Yes | QMessageBox.No, QMessageBox.No)
        if reply == QMessageBox.Yes:
            self.history = []
            self.save_history()
            self.update_history_list()
    
    def load_recent_files(self):
        try:
            with open('recent_files.json', 'r', encoding='utf-8') as f:
                self.recent_files = json.load(f)
                self.update_recent_list()
        except FileNotFoundError:
            self.recent_files = []
    
    def save_recent_files(self):
        with open('recent_files.json', 'w', encoding='utf-8') as f:
            json.dump(self.recent_files, f, ensure_ascii=False, indent=2)
    
    def update_recent_list(self):
        self.recent_list.clear()
        for file_path in reversed(self.recent_files):
            if os.path.exists(file_path):
                file_name = os.path.basename(file_path)
                file_size = os.path.getsize(file_path)
                file_size_str = f"{file_size/1024:.1f}KB" if file_size < 1024*1024 else f"{file_size/1024/1024:.1f}MB"
                self.recent_list.addItem(f"📄 {file_name}\n📊 {file_size_str}")
    
    def add_to_recent(self, file_path):
        if file_path in self.recent_files:
            self.recent_files.remove(file_path)
        self.recent_files.insert(0, file_path)
        # 限制最近文件数量
        if len(self.recent_files) > 10:
            self.recent_files = self.recent_files[:10]
        self.save_recent_files()
        self.update_recent_list()
    
    def load_recent_file(self, item):
        index = self.recent_list.row(item)
        file_path = self.recent_files[-(index + 1)]
        if os.path.exists(file_path):
            self.start_ocr(file_path)
        else:
            QMessageBox.warning(self, "警告", "文件不存在")
            self.recent_files.remove(file_path)
            self.save_recent_files()
            self.update_recent_list()
    
    def show_recent_context_menu(self, position):
        menu = QMenu()
        delete_action = menu.addAction("删除")
        clear_all_action = menu.addAction("清空所有")
        
        item = self.recent_list.itemAt(position)
        if item:
            action = menu.exec_(self.recent_list.mapToGlobal(position))
            if action == delete_action:
                self.delete_recent_item(item)
            elif action == clear_all_action:
                self.clear_all_recent()
        else:
            action = menu.exec_(self.recent_list.mapToGlobal(position))
            if action == clear_all_action:
                self.clear_all_recent()
    
    def delete_recent_item(self, item):
        index = self.recent_list.row(item)
        del self.recent_files[-(index + 1)]
        self.save_recent_files()
        self.update_recent_list()
    
    def clear_all_recent(self):
        reply = QMessageBox.question(self, '确认', '确定要清空所有最近文件吗？',
                                   QMessageBox.Yes | QMessageBox.No, QMessageBox.No)
        if reply == QMessageBox.Yes:
            self.recent_files = []
            self.save_recent_files()
            self.update_recent_list()
    
    def show_cache_settings(self):
        """显示缓存设置对话框"""
        dialog = CacheSettingsDialog(self)
        dialog.exec_()

    def show_baidu_api_settings(self):
        dialog = BaiduAPISettingsDialog(self)
        dialog.exec_()

    def update_font_size(self, size):
        """更新字体大小"""
        # 更新日志窗口字体
        self.log_text.setStyleSheet(f"""
            QTextEdit {{
                background-color: #f0f0f0;
                font-family: 'Consolas', 'Courier New', monospace;
                font-size: {size}px;
                line-height: 1.5;
            }}
        """)
        
        # 更新结果窗口字体
        self.result_text.setStyleSheet(f"""
            QTextEdit {{
                font-family: 'Consolas', 'Courier New', monospace;
                font-size: {size}px;
                line-height: 1.5;
            }}
        """)
        
        # 保存字体大小设置
        self.settings.setValue("font_size", size)

    def batch_export(self):
        """批量导出功能"""
        # 选择多个PDF文件
        files, _ = QFileDialog.getOpenFileNames(
            self,
            "选择PDF文件",
            "",
            "PDF文件 (*.pdf)"
        )
        
        if not files:
            return
            
        # 选择导出格式
        format_dialog = QDialog(self)
        format_dialog.setWindowTitle("选择导出格式")
        layout = QVBoxLayout()
        
        format_combo = QComboBox()
        format_combo.addItems(["TXT", "Word", "PDF"])
        layout.addWidget(QLabel("导出格式:"))
        layout.addWidget(format_combo)
        
        # 选择输出目录
        output_dir = QLineEdit()
        output_dir.setReadOnly(True)
        browse_button = QPushButton("浏览")
        browse_button.clicked.connect(lambda: self._select_output_dir(output_dir))
        
        dir_layout = QHBoxLayout()
        dir_layout.addWidget(QLabel("输出目录:"))
        dir_layout.addWidget(output_dir)
        dir_layout.addWidget(browse_button)
        layout.addLayout(dir_layout)
        
        # 添加按钮
        buttons = QDialogButtonBox(QDialogButtonBox.Ok | QDialogButtonBox.Cancel)
        buttons.accepted.connect(format_dialog.accept)
        buttons.rejected.connect(format_dialog.reject)
        layout.addWidget(buttons)
        
        format_dialog.setLayout(layout)
        
        if format_dialog.exec_() != QDialog.Accepted:
            return
            
        output_dir = output_dir.text()
        if not output_dir:
            QMessageBox.warning(self, "警告", "请选择输出目录")
            return
            
        # 开始批量处理
        self.batch_progress.setVisible(True)
        self.batch_progress.setMaximum(len(files))
        self.batch_progress.setValue(0)
        
        for i, file in enumerate(files):
            try:
                # 处理PDF
                self.current_pdf_path = file
                self.start_ocr(file)
                
                # 等待处理完成
                while self.progress_bar.isVisible():
                    QApplication.processEvents()
                    
                # 获取结果
                result = self.result_text.toPlainText()
                
                # 导出文件
                output_file = os.path.join(
                    output_dir,
                    f"{os.path.splitext(os.path.basename(file))[0]}.{format_combo.currentText().lower()}"
                )
                
                if format_combo.currentText() == "TXT":
                    with open(output_file, 'w', encoding='utf-8') as f:
                        f.write(result)
                elif format_combo.currentText() == "Word":
                    from docx import Document
                    doc = Document()
                    doc.add_paragraph(result)
                    doc.save(output_file)
                elif format_combo.currentText() == "PDF":
                    from reportlab.pdfgen import canvas
                    from reportlab.lib.pagesizes import letter
                    c = canvas.Canvas(output_file, pagesize=letter)
                    textobject = c.beginText()
                    textobject.setTextOrigin(50, 750)
                    textobject.setFont("Helvetica", 12)
                    
                    for line in result.split('\n'):
                        textobject.textLine(line)
                        
                    c.drawText(textobject)
                    c.save()
                    
                self.batch_progress.setValue(i + 1)
                QApplication.processEvents()
                
            except Exception as e:
                self.log_text.append(f"处理文件 {file} 时出错: {str(e)}")
                
        self.batch_progress.setVisible(False)
        QMessageBox.information(self, "完成", "批量导出完成")

    def _select_output_dir(self, output_dir_edit):
        """选择输出目录"""
        dir_path = QFileDialog.getExistingDirectory(
            self,
            "选择输出目录"
        )
        if dir_path:
            output_dir_edit.setText(dir_path)

    def _update_preview(self, text, stats):
        """更新预览和统计信息"""
        # 获取当前处理的图像
        if hasattr(self.current_worker, 'current_image_data'):
            qimage = self.current_worker.current_image_data
        else:
            qimage = QImage()
            
        # 更新预览
        self.preview_widget.update_preview(qimage, text, stats)
        
        # 更新统计信息
        self.stats_labels['pages'].setText(f"页数: {stats['pages']}")
        self.stats_labels['lines'].setText(f"行数: {stats['lines']}")
        self.stats_labels['words'].setText(f"字数: {stats['words']}")
        self.stats_labels['chars'].setText(f"字符数: {stats['chars']}")
        self.stats_labels['confidence'].setText(f"置信度: {stats['confidence']}")

class ProofreadDialog(QDialog):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.setWindowTitle("文本校对")
        self.setModal(True)
        
        layout = QVBoxLayout()
        
        # 创建文本编辑区域
        self.text_edit = QTextEdit()
        self.text_edit.setMinimumSize(600, 400)
        layout.addWidget(self.text_edit)
        
        # 添加按钮
        buttons = QDialogButtonBox(QDialogButtonBox.Ok | QDialogButtonBox.Cancel)
        buttons.accepted.connect(self.accept)
        buttons.rejected.connect(self.reject)
        layout.addWidget(buttons)
        
        self.setLayout(layout)
        
    def set_text(self, text):
        """设置校对文本"""
        self.text_edit.setText(text)
        
    def get_text(self):
        """获取校对后的文本"""
        return self.text_edit.toPlainText()

class PreviewWidget(QWidget):
    def __init__(self, parent=None):
        super().__init__(parent)
        layout = QVBoxLayout()
        
        # 创建分割器
        splitter = QSplitter(Qt.Vertical)
        
        # 图像预览区域
        self.image_label = QLabel()
        self.image_label.setAlignment(Qt.AlignCenter)
        self.image_label.setMinimumSize(400, 300)
        self.image_label.setStyleSheet("""
            QLabel { 
                background-color: white;
                border: 1px solid #cccccc;
            }
        """)
        
        # 文本预览区域
        self.text_edit = QTextEdit()
        self.text_edit.setReadOnly(True)
        self.text_edit.setMinimumHeight(200)
        self.text_edit.setStyleSheet("""
            QTextEdit {
                background-color: #f0f0f0;
                font-family: 'Consolas', 'Courier New', monospace;
                border: 1px solid #cccccc;
            }
        """)
        
        # 添加组件到分割器
        splitter.addWidget(self.image_label)
        splitter.addWidget(self.text_edit)
        
        # 设置分割器比例
        splitter.setSizes([400, 200])
        
        layout.addWidget(splitter)
        self.setLayout(layout)
        
    def update_preview(self, image, text, stats):
        """更新预览内容"""
        # 显示图像
        if image and not image.isNull():
            # 调整图像大小以适应预览区域
            scaled_image = image.scaled(
                self.image_label.size(),
                Qt.KeepAspectRatio,
                Qt.SmoothTransformation
            )
            self.image_label.setPixmap(QPixmap.fromImage(scaled_image))
        else:
            self.image_label.clear()
            self.image_label.setText("正在加载图像...")
            
        # 显示文本
        if text:
            self.text_edit.setText(text)
        else:
            self.text_edit.clear()
            
        # 更新状态栏
        self.parent().statusBar().showMessage(
            f"正在处理第 {stats['pages']} 页 | "
            f"已识别: {stats['words']}字 | "
            f"置信度: {stats['confidence']}"
        )

if __name__ == '__main__':
    app = QApplication(sys.argv)
    window = MainWindow()
    window.show()
    sys.exit(app.exec_()) 